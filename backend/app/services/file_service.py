"""File parsing service — extracts plain text from .txt, .doc/.docx, .pdf.

A PDF is not always text. Scans, photographed notes and anything printed to
PDF from an image carry no text layer at all, and the extractors below return
an empty string for them — which used to travel all the way to the model as an
empty note. Those files are rendered to page images instead and read by a
vision-capable model; see ``parse_document``.
"""

from __future__ import annotations

import base64
import io
import logging
from dataclasses import dataclass, field

from app.models.schemas import DocumentImage

logger = logging.getLogger("mediextract.services.file")

# Below this many characters per page, a PDF is treated as a scan. Not zero:
# a scanned note is often produced by a tool that stamps a header, a page
# number or a footer into the text layer, and a dozen characters of furniture
# is not a note. Anything with real prose in it clears this comfortably.
MIN_TEXT_CHARS_PER_PAGE = 24


@dataclass
class ParsedDocument:
    """The result of reading one uploaded file."""

    text: str = ""
    images: list[DocumentImage] = field(default_factory=list)
    page_count: int = 0
    warning: str = ""

    @property
    def has_content(self) -> bool:
        return bool(self.text.strip() or self.images)


class FileService:
    """Stateless service for extracting text from uploaded documents."""

    def __init__(
        self,
        *,
        render_dpi: int = 150,
        max_render_pages: int = 8,
        jpeg_quality: int = 80,
    ) -> None:
        self.render_dpi = render_dpi
        self.max_render_pages = max_render_pages
        self.jpeg_quality = jpeg_quality

    # ── Public API ──

    def parse_document(self, content: bytes, extension: str) -> ParsedDocument:
        """Read a file into text, or — for a PDF that has none — page images."""
        ext = extension.lower().lstrip(".")
        if ext != "pdf":
            return ParsedDocument(text=self.extract_text(content, ext))

        text, page_count = self._parse_pdf(content)
        if self._is_scanned(text, page_count):
            images = self._render_pdf_pages(content)
            if images:
                skipped = max(page_count - len(images), 0)
                warning = (
                    f"No text layer — reading {len(images)} page"
                    f"{'' if len(images) == 1 else 's'} as images"
                )
                if skipped:
                    warning += f" (first {len(images)} of {page_count})"
                logger.info(
                    "PDF has no usable text layer (%d chars over %d pages) — "
                    "rendered %d page(s) for a vision model",
                    len(text.strip()),
                    page_count,
                    len(images),
                )
                # The stray text is dropped on purpose: a page-number stamp
                # from a scanner is not part of the note, and passing it
                # alongside the images only invites the model to anchor on it.
                return ParsedDocument(
                    images=images, page_count=page_count, warning=warning
                )
            logger.warning("PDF has no text layer and could not be rendered to images")

        return ParsedDocument(text=text, page_count=page_count)

    def extract_text(self, content: bytes, extension: str) -> str:
        """Dispatch to the correct parser based on file extension."""
        ext = extension.lower().lstrip(".")
        match ext:
            case "txt":
                return self._parse_txt(content)
            case "pdf":
                return self._parse_pdf(content)[0]
            case "doc" | "docx":
                return self._parse_docx(content)
            case _:
                raise ValueError(f"Unsupported extension: .{ext}")

    # ── Parsers ──

    @staticmethod
    def _parse_txt(content: bytes) -> str:
        """Plain text — decode with fallback."""
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                return content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue
        return content.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_pdf(content: bytes) -> tuple[str, int]:
        """Extract the text layer of a PDF, and count its pages.

        The page count comes back because it is what says whether the text
        that *was* found amounts to a document or to a scanner's furniture —
        a decision the caller cannot make from the text alone.
        """
        try:
            import pdfplumber
            text_parts: list[str] = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            if text_parts:
                return "\n\n".join(text_parts), page_count
        except Exception:
            logger.debug("pdfplumber failed, falling back to PyMuPDF")

        # Fallback: PyMuPDF
        import fitz  # pymupdf
        doc = fitz.open(stream=content, filetype="pdf")
        pages = [page.get_text() for page in doc]
        page_count = doc.page_count
        doc.close()
        return "\n\n".join(pages), page_count

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        """Extract text from .doc/.docx using python-docx."""
        from docx import Document

        doc = Document(io.BytesIO(content))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())

    # ── Scans ──

    @staticmethod
    def _is_scanned(text: str, page_count: int) -> bool:
        """Whether a PDF's text layer is too thin to be the document itself."""
        return len(text.strip()) < MIN_TEXT_CHARS_PER_PAGE * max(page_count, 1)

    def _render_pdf_pages(self, content: bytes) -> list[DocumentImage]:
        """Render the first pages of a PDF to JPEG images.

        Capped at ``max_render_pages``: each page is a few hundred KB of
        base64 that has to survive two JSON round trips and a model's context
        window, and a fifty-page scan is a different feature than this one.
        """
        try:
            import fitz  # pymupdf
        except ImportError:
            logger.error("pymupdf is not installed — cannot render scanned PDFs")
            return []

        images: list[DocumentImage] = []
        try:
            doc = fitz.open(stream=content, filetype="pdf")
        except Exception:
            logger.exception("Could not open PDF for rendering")
            return []

        try:
            for index, page in enumerate(doc):
                if index >= self.max_render_pages:
                    logger.info(
                        "PDF has %d pages — rendering the first %d",
                        doc.page_count,
                        self.max_render_pages,
                    )
                    break
                try:
                    pixmap = page.get_pixmap(dpi=self.render_dpi)
                    try:
                        data = pixmap.tobytes("jpeg", jpg_quality=self.jpeg_quality)
                        mime = "image/jpeg"
                    except Exception:
                        # Older PyMuPDF builds, or a pixmap with an alpha
                        # channel, refuse JPEG. PNG is always available.
                        data = pixmap.tobytes("png")
                        mime = "image/png"
                except Exception:
                    logger.exception("Could not render page %d", index + 1)
                    continue
                images.append(
                    DocumentImage(
                        mime_type=mime,
                        data=base64.b64encode(data).decode("ascii"),
                        page=index + 1,
                    )
                )
        finally:
            doc.close()

        return images
